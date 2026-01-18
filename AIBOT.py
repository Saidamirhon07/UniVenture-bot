# AIBOT.py
# UniVentureAI — Telegram bot with GLOBAL per-topic RAG + metadata separation (qa/evaluation)
# + eval follow-ups (apply feedback) + eval Q&A (any follow-up question about the evaluated text)
# + embedded tools + Application Plan & School Finder
# + analytics + admin locks + backup + health + robust command parsing + UUID doc IDs (no reteach bugs)

import os
os.environ['TZ'] = 'UTC'  # Set timezone to UTC

from telegram import Update, KeyboardButton, ReplyKeyboardMarkup
from telegram.constants import ChatAction
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters,
)
from dotenv import load_dotenv
import chromadb
from chromadb.utils import embedding_functions
import os, io, nest_asyncio, logging, json, base64, uuid

# -------- File extraction deps --------
from pdfminer.high_level import extract_text
from docx import Document as DocxDocument

# -------- Web page extraction (for teachlink) --------
import trafilatura

import threading
from http.server import HTTPServer, BaseHTTPRequestHandler

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
load_dotenv()
nest_asyncio.apply()

TELEGRAM_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not TELEGRAM_TOKEN:
    raise RuntimeError("Missing TELEGRAM_BOT_TOKEN in environment.")
if not OPENAI_API_KEY:
    raise RuntimeError("Missing OPENAI_API_KEY in environment.")

# =========================
# OpenAI SDK compatibility
# =========================
USE_NEW_OPENAI = False
_client = None
try:
    # openai>=1.x
    from openai import OpenAI  # type: ignore

    _client = OpenAI(api_key=OPENAI_API_KEY)
    USE_NEW_OPENAI = True
    logging.info("Using OpenAI SDK v1.x+")
except Exception as e:
    # openai<=0.28.x
    import openai  # type: ignore

    openai.api_key = OPENAI_API_KEY
    USE_NEW_OPENAI = False
    logging.info("Using OpenAI SDK v0.28.x or earlier")

def openai_chat(
    model: str,
    messages: list,
    temperature: float = 0.4,
    max_tokens: int | None = None,
) -> str:
    """Unified chat completion across old/new OpenAI python SDKs."""
    try:
        if USE_NEW_OPENAI:
            resp = _client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            return (resp.choices[0].message.content or "").strip()
        else:
            import openai as _openai  # type: ignore

            resp = _openai.ChatCompletion.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            return (resp["choices"][0]["message"]["content"] or "").strip()
    except Exception as e:
        logging.exception(f"OpenAI error in model {model}")
        return f"Error: {str(e)[:200]}"

# -------- Admin config --------
ADMIN_IDS = {
    886181760,  # TODO: replace with YOUR Telegram user ID (from @userinfobot)
}

def require_admin(update: Update) -> bool:
    user = update.effective_user
    return bool(user and user.id in ADMIN_IDS)

# -------- Data / storage config (for Railway/Render volume etc.) --------
DATA_DIR = os.getenv("DATA_DIR", "./data")
os.makedirs(DATA_DIR, exist_ok=True)

CHROMA_PATH = os.getenv("CHROMA_PATH", os.path.join(DATA_DIR, "chroma_store"))
COLLECTION_PREFIX = os.getenv("CHROMA_COLLECTION_PREFIX", "global")

# -------- Persistent Chroma (GLOBAL per-topic collections) --------
try:
    chroma = chromadb.PersistentClient(path=CHROMA_PATH)
    logging.info(f"ChromaDB initialized at {CHROMA_PATH}")
except Exception as e:
    logging.error(f"Failed to initialize ChromaDB: {e}")
    raise

emb_fn = embedding_functions.OpenAIEmbeddingFunction(
    api_key=OPENAI_API_KEY,
    model_name="text-embedding-3-small",
)

# -------- Analytics storage --------
STATS_FILE = os.path.join(DATA_DIR, "analytics.json")

def _default_stats():
    return {
        "users": [],
        "messages_total": 0,
        "messages_per_user": {},
        "topic_counts": {},
        "eval_counts": {},
    }

def load_stats():
    if not os.path.exists(STATS_FILE):
        return _default_stats()
    try:
        with open(STATS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        logging.warning(f"Could not load analytics: {e}")
        return _default_stats()

    base = _default_stats()
    base.update({k: data.get(k, v) for k, v in base.items()})
    return base

def save_stats(stats):
    try:
        with open(STATS_FILE, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logging.warning(f"Could not save analytics: {e}")

def record_event(user_id, topic: str, kind: str = "message"):
    stats = load_stats()
    uid = str(user_id)

    if uid not in stats["users"]:
        stats["users"].append(uid)

    stats["messages_total"] = stats.get("messages_total", 0) + 1

    mpu = stats.setdefault("messages_per_user", {})
    mpu[uid] = mpu.get(uid, 0) + 1

    if topic:
        tc = stats.setdefault("topic_counts", {})
        tc[topic] = tc.get(topic, 0) + 1

    if kind == "eval" and topic:
        ec = stats.setdefault("eval_counts", {})
        ec[topic] = ec.get(topic, 0) + 1

    save_stats(stats)

# -------- Main menu buttons --------
BTN_ESSAY = "📝 Essays"
BTN_EC = "🎯 Extracurricular activities"
BTN_REC = "✉️ Recommendation Letters"
BTN_SAT = "📈 SAT"
BTN_IELTS = "🗣️ IELTS"
BTN_PORT = "🖼️ Portfolio Check"
BTN_PLAN_MAIN = "📅 Application Plan"
BTN_SF_MAIN = "🏫 School Finder"

# Boost tools (cross-topic)
BTN_TOOLS = "🚀 Boost Tools"
BTN_PROGRESS = "📊 My Progress"
BTN_INSIDER = "🤫 Insider Tips"
BTN_POWERWORDS = "⚡ Power Words"
BTN_PREDICT = "🎯 Predict My Chances"
BTN_WOWFACTOR = "🔍 Find Wow Factor"

# Evaluation sub-buttons
BTN_PS_EVAL = "✅ Personal Statement Evaluation"
BTN_SUPP_EVAL = "✅ Supplemental Essay Evaluation"
BTN_EC_EVAL = "✅ Extracurricular Evaluation"
BTN_REC_EVAL = "✅ Rec Letter Evaluation"
BTN_IW_EVAL = "✅ Writing Evaluation"
BTN_PORT_EVAL = "✅ Portfolio Evaluation"

# Essays sub-buttons
BTN_ESSAY_PS = "📝 Personal Statement"
BTN_ESSAY_SUPP = "📝 Supplemental Essays"

# Extra tools (as buttons inside menus)
BTN_EC_PROGRAMS = "🌍 Top Programs & Opportunities"
BTN_BRAINSTORM = "🧠 Brainstorm ideas"
BTN_REWRITE = "✍️ Rewrite my text"
BTN_REC_PACKET = "📄 Rec Letter Packet"
BTN_PORTFOLIO_IDEAS = "💡 Portfolio Ideas"

# SAT sub-buttons
BTN_SAT_MATH = "📐 SAT Math"
BTN_SAT_ENGLISH = "📚 SAT English"

# IELTS sub-buttons
BTN_IELTS_READING = "📖 IELTS Reading"
BTN_IELTS_LISTENING = "👂 IELTS Listening"
BTN_IELTS_WRITING = "✍️ IELTS Writing"
BTN_IELTS_SPEAKING = "🗣️ IELTS Speaking"

# Back button
BTN_BACK = "⬅️ Back"

# -------- Keyboards --------
def main_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_ESSAY), KeyboardButton(BTN_EC), KeyboardButton(BTN_REC)],
            [KeyboardButton(BTN_SAT), KeyboardButton(BTN_IELTS), KeyboardButton(BTN_PORT)],
            [KeyboardButton(BTN_PLAN_MAIN), KeyboardButton(BTN_SF_MAIN), KeyboardButton(BTN_TOOLS)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def tools_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PROGRESS), KeyboardButton(BTN_INSIDER), KeyboardButton(BTN_POWERWORDS)],
            [KeyboardButton(BTN_PREDICT), KeyboardButton(BTN_WOWFACTOR)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def essay_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_ESSAY_PS), KeyboardButton(BTN_ESSAY_SUPP)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def essay_ps_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PS_EVAL)],
            [KeyboardButton(BTN_BRAINSTORM)],
            [KeyboardButton(BTN_REWRITE)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def essay_supp_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_SUPP_EVAL)],
            [KeyboardButton(BTN_BRAINSTORM)],
            [KeyboardButton(BTN_REWRITE)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def ec_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_EC_EVAL)],
            [KeyboardButton(BTN_EC_PROGRAMS)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def rec_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_REC_EVAL)],
            [KeyboardButton(BTN_REC_PACKET)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def sat_menu_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_SAT_MATH), KeyboardButton(BTN_SAT_ENGLISH)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def ielts_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_IELTS_READING), KeyboardButton(BTN_IELTS_LISTENING)],
            [KeyboardButton(BTN_IELTS_WRITING), KeyboardButton(BTN_IELTS_SPEAKING)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def ielts_writing_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_IW_EVAL)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def portfolio_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton(BTN_PORT_EVAL)],
            [KeyboardButton(BTN_PORTFOLIO_IDEAS)],
            [KeyboardButton(BTN_BACK)],
        ],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def plan_keyboard():
    return ReplyKeyboardMarkup(
        [[KeyboardButton(BTN_BACK)]],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def schoolfinder_keyboard():
    return ReplyKeyboardMarkup(
        [[KeyboardButton(BTN_BACK)]],
        resize_keyboard=True,
        one_time_keyboard=False,
    )

def is_back_message(text: str) -> bool:
    t = (text or "").strip()
    return t == BTN_BACK or t.lower() == "back"

# -------- Topic mapping & helpers --------
TOPIC_KEYS = {
    BTN_ESSAY_PS: "essays_personal",
    BTN_ESSAY_SUPP: "essays_supplemental",
    BTN_EC: "extracurriculars",
    BTN_REC: "recommendations",
    BTN_PORT: "portfolio",
    BTN_SAT_MATH: "sat_math",
    BTN_SAT_ENGLISH: "sat_english",
    BTN_IELTS_READING: "ielts_reading",
    BTN_IELTS_LISTENING: "ielts_listening",
    BTN_IELTS_WRITING: "ielts_writing",
    BTN_IELTS_SPEAKING: "ielts_speaking",
}

EVAL_TOPICS = {
    "essays_personal",
    "essays_supplemental",
    "recommendations",
    "portfolio",
    "extracurriculars",
}

FRIENDLY_TOPIC_NAMES = {
    "essays_personal": "Personal Statement",
    "essays_supplemental": "Supplemental Essays",
    "extracurriculars": "Extracurricular Activities",
    "recommendations": "Recommendation Letters",
    "portfolio": "Portfolio",
    "sat_math": "SAT Math",
    "sat_english": "SAT English",
    "ielts_reading": "IELTS Reading",
    "ielts_listening": "IELTS Listening",
    "ielts_writing": "IELTS Writing",
    "ielts_speaking": "IELTS Speaking",
    "application_plan": "Application Plan",
    "school_finder": "School Finder",
    "general": "General admissions help",
}

DEFAULT_TOPIC = "general"

# -------- Image mapping --------
IMAGE_FILES = {
    "welcome": "images/welcome.png",
    "essays_main": "images/essays_main.png",
    "essays_personal": "images/essays_personal.png",
    "essays_supplemental": "images/essays_supplemental.png",
    "extracurriculars": "images/extracurriculars.png",
    "recommendations": "images/recommendations.png",
    "portfolio": "images/portfolio.png",
    "sat_main": "images/sat_main.png",
    "sat_math": "images/sat_math.png",
    "sat_english": "images/sat_english.png",
    "ielts_main": "images/ielts_main.png",
    "ielts_reading": "images/ielts_reading.png",
    "ielts_listening": "images/ielts_listening.png",
    "ielts_writing": "images/ielts_writing.png",
    "ielts_speaking": "images/ielts_speaking.png",
    "plan_main": "images/plan_main.png",
    "schoolfinder_main": "images/schoolfinder_main.png",
}

def get_current_topic(context: ContextTypes.DEFAULT_TYPE) -> str:
    return context.user_data.get("topic", DEFAULT_TOPIC)

def track_topic(context: ContextTypes.DEFAULT_TYPE, topic: str):
    """Track which topic menus the user has opened (per-user, stored in user_data)."""
    if not topic:
        return
    seen = set(context.user_data.get("topics_seen", []))
    seen.add(topic)
    context.user_data["topics_seen"] = sorted(seen)

def track_tool_use(context: ContextTypes.DEFAULT_TYPE, tool: str):
    if not tool:
        return
    used = set(context.user_data.get("tools_used", []))
    used.add(tool)
    context.user_data["tools_used"] = sorted(used)

def get_collection(chat_id: int, topic: str):
    # GLOBAL per-topic collections shared by all users.
    # chat_id kept for backwards compatibility but ignored.
    collection_name = f"{COLLECTION_PREFIX}_{topic}"
    try:
        return chroma.get_or_create_collection(
            name=collection_name,
            embedding_function=emb_fn,
        )
    except Exception as e:
        logging.error(f"Error getting collection {collection_name}: {e}")
        raise

def new_doc_id(topic: str, tag: str = "") -> str:
    """Always-generate unique IDs for Chroma (prevents unlearn→reteach collisions)."""
    u = uuid.uuid4().hex
    return f"{topic}_{tag}_{u}" if tag else f"{topic}_{u}"

def _chunk(text: str, max_chars=1000, overlap=150):
    text = text or ""
    if not text.strip():
        return []
    chunks, i = [], 0
    while i < len(text):
        end = min(len(text), i + max_chars)
        chunks.append(text[i:end])
        if end == len(text):
            break
        i = max(0, end - overlap)
    return chunks

def extract_command_text(update: Update) -> str:
    msg = update.message
    if not msg:
        return ""
    if msg.text:
        return msg.text.strip()
    if msg.caption:
        return msg.caption.strip()
    return ""

def strip_command(text: str, command: str) -> str:
    """
    Remove the first token (/command or /command@BotName) and return the rest.
    Fixes bugs where titles accidentally include '@BotName'.
    """
    t = (text or "").strip()
    if not t:
        return ""
    parts = t.split(maxsplit=1)
    if len(parts) < 1:
        return ""
    
    first = parts[0]
    rest = parts[1] if len(parts) > 1 else ""
    
    # Check if first part is the command (with or without bot username)
    if first.startswith(f"/{command}"):
        return rest.strip()
    
    # If not, return original text (in case user typed something else)
    return t

def is_caption_command(caption: str, cmd: str) -> bool:
    cap = (caption or "").strip()
    if not cap:
        return False
    first = cap.split(maxsplit=1)[0]
    return first.startswith(f"/{cmd}")

async def show_typing(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    try:
        await context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
    except Exception as e:
        logging.warning(f"Could not send typing action: {e}")

def sanitize_output(text: str) -> str:
    if not text:
        return text
    return text.replace("—", " - ")

def safe_text_for_embedding(text: str) -> str:
    if not text:
        return text
    return text.encode("utf-8", "ignore").decode("utf-8")

def _truncate_for_storage(text: str, max_chars: int = 12000) -> str:
    if not text:
        return ""
    t = text.strip()
    return t[:max_chars]

async def send_long(update: Update, text: str):
    MAX_LEN = 4000
    if not text:
        return
    text = sanitize_output(text)
    for i in range(0, len(text), MAX_LEN):
        try:
            await update.message.reply_text(text[i : i + MAX_LEN])
        except Exception as e:
            logging.error(f"Failed to send message part: {e}")

async def send_with_image(
    update: Update,
    caption: str,
    reply_markup=None,
    image_key: str | None = None,
):
    if image_key and image_key in IMAGE_FILES:
        path = IMAGE_FILES[image_key]
        if os.path.exists(path):
            try:
                with open(path, "rb") as f:
                    await update.message.reply_photo(
                        photo=f,
                        caption=caption,
                        reply_markup=reply_markup,
                    )
                return
            except Exception as e:
                logging.warning(f"Failed to send image {path}: {e}")
        else:
            logging.warning(f"Image file not found: {path}")

    await update.message.reply_text(caption, reply_markup=reply_markup)

# -------- Vision helper --------
def extract_text_from_image_bytes(image_bytes: bytes) -> str:
    try:
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        data_url = f"data:image/jpeg;base64,{b64}"

        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract all readable text from this image. "
                            "Return ONLY the plain text, no extra comments."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ]

        out = openai_chat(model="gpt-4o-mini", messages=messages, temperature=0.0)
        return out or ""
    except Exception as e:
        logging.error(f"Error extracting text from image: {e}")
        return ""

# ---------- EVAL FOLLOW-UP HELPERS ----------
FOLLOWUP_TRIGGERS = [
    "next step",
    "do it",
    "do this",
    "apply",
    "apply this",
    "apply the feedback",
    "rewrite",
    "rewrite the ending",
    "rewrite the conclusion",
    "improve",
    "fix",
    "edit",
    "revise",
    "expand",
    "continue",
    "make it smoother",
    "make transitions",
    "stronger conclusion",
    "deeper reflection",
]

def is_followup_intent(q: str) -> bool:
    s = (q or "").lower()
    return any(t in s for t in FOLLOWUP_TRIGGERS)

def looks_like_submission(q: str) -> bool:
    t = (q or "").strip()
    if len(t) >= 600:
        return True
    if len(t) >= 250 and t.count("\n") >= 6:
        return True
    return False

def clear_eval_context(context: ContextTypes.DEFAULT_TYPE):
    for k in [
        "last_eval_topic",
        "last_eval_text",
        "last_eval_text_original",
        "last_eval_feedback",
    ]:
        context.user_data.pop(k, None)
    context.user_data.pop("eval_active", None)

def set_eval_context(
    context: ContextTypes.DEFAULT_TYPE,
    topic: str,
    student_text: str,
    feedback: str,
):
    context.user_data["last_eval_topic"] = topic
    context.user_data["last_eval_text_original"] = student_text
    context.user_data["last_eval_text"] = student_text  # updated on rewrites
    context.user_data["last_eval_feedback"] = feedback
    context.user_data["eval_active"] = True

def _pretty_topic_for_eval(topic: str) -> str:
    return {
        "essays_personal": "Personal Statement essay",
        "essays_supplemental": "Supplemental essay",
        "recommendations": "Recommendation letter",
        "portfolio": "Portfolio description",
        "extracurriculars": "Extracurricular activities description",
        "ielts_writing": "IELTS Writing answer",
    }.get(topic, "document")

def _sys_role_for_eval(topic: str) -> str:
    if topic in {"essays_personal", "essays_supplemental"}:
        return (
            "You are an expert college admissions essay coach. "
            "Improve the student's essay using the prior feedback. "
            "Focus on clarity, structure, voice, authenticity, reflection, and impact."
        )
    if topic == "recommendations":
        return (
            "You are an expert on college recommendation letters. "
            "Improve the letter using the prior feedback. "
            "Focus on specificity, credibility, depth of insight, and support for the student."
        )
    if topic == "extracurriculars":
        return (
            "You are an expert on extracurricular strategy for college applications. "
            "Improve the EC descriptions using the prior feedback. "
            "Focus on impact, leadership, continuity, clarity, and strong phrasing."
        )
    if topic == "ielts_writing":
        return (
            "You are an experienced IELTS Writing examiner. "
            "Improve the student's writing using the prior feedback. "
            "Focus on Task Response, Coherence and Cohesion, Lexical Resource, and Grammar."
        )
    return (
        "You are an expert college portfolio reviewer. "
        "Improve the portfolio description using the prior feedback. "
        "Focus on coherence, originality, technical quality, and fit for selective colleges."
    )

# ---------- Handlers ----------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    clear_eval_context(context)
    context.user_data["topic"] = DEFAULT_TOPIC
    context.user_data["pending_feature"] = None
    context.user_data.pop("in_tools", None)
    user = update.effective_user
    record_event(user.id, "start", kind="start")

    await send_with_image(
        update,
        "Hi! I'm your coached AI 🤖\nChoose a topic or ask a question.",
        reply_markup=main_menu_keyboard(),
        image_key="welcome",
    )

# ---------- TEACH (Q&A sources) ----------
async def teach(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logging.info("🔥 /teach RECEIVED")

    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach global sources.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teach")

    raw = strip_command(extract_command_text(update), "teach")
    logging.info(f"Raw teach input: '{raw}'")
    
    if not raw:
        await update.message.reply_text(
            "Use format:\n/teach <title> | <content>\n\n" f"Current topic: {topic}"
        )
        return
        
    if "|" not in raw:
        await update.message.reply_text(
            "Use format:\n/teach <title> | <content>\n\n" f"Current topic: {topic}"
        )
        return

    try:
        title, content = [p.strip() for p in raw.split("|", 1)]
    except ValueError:
        await update.message.reply_text("❌ Invalid format. Use: /teach <title> | <content>")
        return
        
    if not title or not content:
        await update.message.reply_text("❌ Title or content is empty. Use: /teach <title> | <content>")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": title, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' already exists in topic: {topic}. "
            "Use /unlearn '{title}' first if you want to replace it."
        )
        return

    doc_id = new_doc_id(topic, "qa")
    try:
        col.add(
            ids=[doc_id],
            metadatas=[{"title": title, "topic": topic, "type": "qa", "source": "manual"}],
            documents=[safe_text_for_embedding(content)],
        )
        logging.info(f"Successfully added document '{title}' to topic '{topic}'")
        await update.message.reply_text(
            f"Learned '{title}' ✅ (topic: {topic}, mode: Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add document to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save: {str(e)[:200]}")

# ---------- TEACH RUBRIC (EVALUATION sources) ----------
async def teachrubric(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach global rubrics.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user

    raw = strip_command(extract_command_text(update), "teachrubric")
    logging.info(f"Raw teachrubric input: '{raw}'")
    
    if not raw:
        await update.message.reply_text(
            "Use format:\n/teachrubric <title> | <rubric / evaluation criteria>"
        )
        return
        
    if "|" not in raw:
        await update.message.reply_text(
            "Use format:\n/teachrubric <title> | <rubric / evaluation criteria>"
        )
        return

    try:
        title, content = [p.strip() for p in raw.split("|", 1)]
    except ValueError:
        await update.message.reply_text("❌ Invalid format. Use: /teachrubric <title> | <rubric>")
        return
        
    if not title or not content:
        await update.message.reply_text("❌ Title or rubric content is empty.")
        return

    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachrubric")

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": title, "type": "evaluation"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' already exists in topic: {topic}. "
            "Use /unlearn '{title}' first if you want to replace it."
        )
        return

    try:
        col.add(
            ids=[new_doc_id(topic, "eval")],
            metadatas=[
                {"title": title, "topic": topic, "type": "evaluation", "source": "manual"}
            ],
            documents=[safe_text_for_embedding(content)],
        )
        logging.info(f"Successfully added rubric '{title}' to topic '{topic}'")
        await update.message.reply_text(
            f"Learned evaluation rubric '{title}' ✅ (topic: {topic}, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add rubric to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save rubric: {str(e)[:200]}")

# ---------- TEACH FILE (Q&A sources) ----------
async def teachfile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach from files.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachfile")

    doc = update.message.document
    if not doc:
        await update.message.reply_text(
            "Attach a PDF or DOCX and write /teachfile in the caption to train me from it."
        )
        return

    await update.message.reply_text(
        "Reading your file and extracting text to learn from it (Q&A)…"
    )

    try:
        tgfile = await doc.get_file()
        file_bytes = await tgfile.download_as_bytearray()
        name = (doc.file_name or "upload").lower()
    except Exception as e:
        await update.message.reply_text(f"Failed to download file: {e}")
        return

    try:
        if name.endswith(".pdf"):
            text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for /teachfile.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    if not text or not text.strip():
        await update.message.reply_text("I couldn't find any readable text in that file.")
        return

    parts = _chunk(text)
    if not parts:
        await update.message.reply_text("Text was too short or could not be chunked.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": name, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{name}' is already learned in topic: {topic}.\n"
            "Use /unlearn <title> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "qa") for _ in range(len(parts))]
        metas = [
            {"title": name, "topic": topic, "part": i, "source": "file", "type": "qa"}
            for i in range(len(parts))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(p) for p in parts])
        logging.info(f"Successfully added file '{name}' with {len(parts)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned from file ✅ ({len(parts)} parts) in topic: {topic} (Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add file to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save file content: {str(e)[:200]}")

# ---------- TEACH FILE EVAL (EVALUATION sources) ----------
async def teachfile_eval(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text(
            "⛔ You are not allowed to teach evaluation rubrics from files."
        )
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachfile_eval")

    doc = update.message.document
    if not doc:
        await update.message.reply_text(
            "Attach a PDF or DOCX and write /teachfile_eval in the caption to teach an evaluation rubric."
        )
        return

    await update.message.reply_text(
        "Reading your rubric file and extracting evaluation criteria…"
    )

    try:
        tgfile = await doc.get_file()
        file_bytes = await tgfile.download_as_bytearray()
        name = (doc.file_name or "upload").lower()
    except Exception as e:
        await update.message.reply_text(f"Failed to download file: {e}")
        return

    try:
        if name.endswith(".pdf"):
            text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for /teachfile_eval.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    if not text or not text.strip():
        await update.message.reply_text("I couldn't find any readable text in that file.")
        return

    parts = _chunk(text)
    if not parts:
        await update.message.reply_text("Text was too short or could not be chunked.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": name, "type": "evaluation"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{name}' is already learned in topic: {topic}.\n"
            "Use /unlearn <title> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "eval") for _ in range(len(parts))]
        metas = [
            {"title": name, "topic": topic, "part": i, "source": "file", "type": "evaluation"}
            for i in range(len(parts))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(p) for p in parts])
        logging.info(f"Successfully added eval file '{name}' with {len(parts)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned evaluation rubric from file ✅ ({len(parts)} parts) in topic: {topic} (scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add eval file to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save rubric content: {str(e)[:200]}")

# ---------- TEACH LINK (Q&A) ----------
async def teachlink(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach from links.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachlink")

    msg_text = extract_command_text(update)
    rest = strip_command(msg_text, "teachlink")
    if not rest:
        await update.message.reply_text("Use: /teachlink <url>")
        return

    url = rest.strip()
    await update.message.reply_text("Fetching content from link and learning from it (Q&A)…")

    try:
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            await update.message.reply_text("Could not fetch the URL. Please check the link.")
            return
        text = trafilatura.extract(downloaded)
    except Exception as e:
        await update.message.reply_text(f"Error while fetching the URL: {e}")
        return

    if not text:
        await update.message.reply_text("Couldn't extract readable text from that link.")
        return

    chunks = _chunk(text)
    if not chunks:
        await update.message.reply_text("The page did not contain enough text to learn from.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": url, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"This link is already learned in topic: {topic}.\nUse /unlearn <url> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "qa") for _ in range(len(chunks))]
        metas = [
            {"title": url, "topic": topic, "part": i, "source": "link", "type": "qa"}
            for i in range(len(chunks))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(c) for c in chunks])
        logging.info(f"Successfully added link '{url}' with {len(chunks)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned from link ✅ ({len(chunks)} parts) in topic: {topic} (Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add link to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save link content: {str(e)[:200]}")

# ---------- TEACH LINK EVAL (EVALUATION sources) ----------
async def teachlink_eval(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text(
            "⛔ You are not allowed to teach evaluation material from links."
        )
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachlink_eval")

    msg_text = extract_command_text(update)
    rest = strip_command(msg_text, "teachlink_eval")
    if not rest:
        await update.message.reply_text("Use: /teachlink_eval <url>")
        return

    url = rest.strip()
    await update.message.reply_text(
        "Fetching content from link and learning it as evaluation / rubric material…"
    )

    try:
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            await update.message.reply_text("Could not fetch the URL. Please check the link.")
            return
        text = trafilatura.extract(downloaded)
    except Exception as e:
        await update.message.reply_text(f"Error while fetching the URL: {e}")
        return

    if not text:
        await update.message.reply_text("Couldn't extract readable text from that link.")
        return

    chunks = _chunk(text)
    if not chunks:
        await update.message.reply_text("The page did not contain enough text to learn from.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": url, "type": "evaluation"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"This link is already learned in topic: {topic}.\nUse /unlearn <url> to remove it first."
        )
        return

    try:
        ids = [new_doc_id(topic, "eval") for _ in range(len(chunks))]
        metas = [
            {"title": url, "topic": topic, "part": i, "source": "link", "type": "evaluation"}
            for i in range(len(chunks))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(c) for c in chunks])
        logging.info(f"Successfully added eval link '{url}' with {len(chunks)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned evaluation material from link ✅ ({len(chunks)} parts) in topic: {topic} (scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add eval link to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save evaluation content: {str(e)[:200]}")

# ---------- TEACH IMAGE (Q&A) ----------
async def teachimage(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to teach from images.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="teachimage")

    photos = update.message.photo or []
    if not photos:
        await update.message.reply_text("Please send a clear image with caption:\n/teachimage <title>")
        return

    caption = (update.message.caption or "").strip()
    title = None
    if is_caption_command(caption, "teachimage"):
        rest = strip_command(caption, "teachimage")
        title = rest.strip() if rest.strip() else None

    largest = photos[-1]
    try:
        tgfile = await largest.get_file()
        if not title:
            title = f"image_{tgfile.file_unique_id}"
    except Exception as e:
        await update.message.reply_text(f"Failed to get image file: {e}")
        return

    await update.message.reply_text(
        f"Reading your image for topic '{topic}' and extracting text to learn from it…"
    )

    try:
        img_bytes = await tgfile.download_as_bytearray()
    except Exception as e:
        await update.message.reply_text(f"Failed to download image: {e}")
        return

    try:
        extracted = extract_text_from_image_bytes(img_bytes)
    except Exception as e:
        await update.message.reply_text(f"Could not extract text from image: {e}")
        return

    if not extracted.strip():
        await update.message.reply_text("I couldn't read any text from that image.")
        return

    parts = _chunk(extracted)
    if not parts:
        await update.message.reply_text("The extracted text was too short to learn from.")
        return

    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return

    try:
        existing = col.get(where={"title": title, "type": "qa"})
    except Exception as e:
        logging.error(f"Error checking existing: {e}")
        existing = {"ids": []}

    if existing and existing.get("ids"):
        await update.message.reply_text(
            f"'{title}' is already learned in topic: {topic}.\nUse /unlearn <title> to remove it first if needed."
        )
        return

    try:
        ids = [new_doc_id(topic, "qa") for _ in range(len(parts))]
        metas = [
            {"title": title, "topic": topic, "part": i, "source": "image", "type": "qa"}
            for i in range(len(parts))
        ]
        col.add(ids=ids, metadatas=metas, documents=[safe_text_for_embedding(p) for p in parts])
        logging.info(f"Successfully added image '{title}' with {len(parts)} parts to topic '{topic}'")
        await update.message.reply_text(
            f"Learned from image '{title}' ✅ ({len(parts)} parts) in topic: {topic} (Q&A, scope: GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Failed to add image to Chroma: {e}")
        await update.message.reply_text(f"❌ Failed to save image content: {str(e)[:200]}")

# ---------- SOURCES ----------
async def sources_all(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logging.info("🔥 /sources_all handler hit")
    try:
        collections = chroma.list_collections()
    except Exception as e:
        await update.message.reply_text(f"Error accessing database: {e}")
        return

    if not collections:
        await update.message.reply_text("No sources stored yet.")
        return

    source_stats = {}
    total_bytes = 0
    total_chunks = 0

    for col_info in collections:
        col_name = col_info.name if hasattr(col_info, "name") else str(col_info)
        if not col_name:
            continue
            
        try:
            col = chroma.get_collection(col_name)
            data = col.get(include=["documents", "metadatas"])
        except Exception as e:
            logging.error(f"Error getting collection {col_name}: {e}")
            continue

        docs = data.get("documents", [])
        metas = data.get("metadatas", [])

        for doc, meta in zip(docs, metas):
            title = (meta or {}).get("title", "Untitled")
            size_bytes = len((doc or "").encode("utf-8"))

            if title not in source_stats:
                source_stats[title] = {"chunks": 0, "bytes": 0}

            source_stats[title]["chunks"] += 1
            source_stats[title]["bytes"] += size_bytes

            total_chunks += 1
            total_bytes += size_bytes

    if not source_stats:
        await update.message.reply_text("No sources found in any collections.")
        return

    lines = []
    for title, stats in sorted(source_stats.items(), key=lambda x: x[1]["bytes"], reverse=True):
        mb = stats["bytes"] / (1024 * 1024)
        lines.append(f"• {title}: {stats['chunks']} chunks, {mb:.2f} MB")

    total_mb = total_bytes / (1024 * 1024)
    msg = (
        "📚 ALL BOT SOURCES (ALL TOPICS, GLOBAL)\n\n"
        f"Total chunks: {total_chunks}\n"
        f"Total text size: {total_mb:.2f} MB\n\n"
        + "\n".join(lines)
    )
    await send_long(update, msg)

async def sources(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logging.info("🔥 /sources handler hit")
    chat_id = update.effective_chat.id
    topic = get_current_topic(context)
    
    try:
        col = get_collection(chat_id, topic)
        data = col.get(include=["metadatas"])
    except Exception as e:
        logging.error(f"Error getting collection: {e}")
        await update.message.reply_text(f"Error accessing database: {e}")
        return
        
    metas = data.get("metadatas") or []

    if not metas:
        await update.message.reply_text(
            f"📌 Active topic: {topic}\nNo sources yet. (Global collection is empty.)"
        )
        return

    qa_titles = []
    eval_titles = []
    for m in metas:
        title = (m or {}).get("title", "Untitled")
        source_type = (m or {}).get("type", "qa")
        if source_type == "evaluation":
            eval_titles.append(title)
        else:
            qa_titles.append(title)

    qa_titles = list(dict.fromkeys(qa_titles))
    eval_titles = list(dict.fromkeys(eval_titles))

    lines = [f"📌 Active topic: {topic} (GLOBAL)\n"]
    if eval_titles:
        lines.append("📘 Evaluation Rubrics:")
        for t in eval_titles:
            lines.append(f"• {t}")
        lines.append("")
    if qa_titles:
        lines.append("📗 Q&A Sources:")
        for t in qa_titles:
            lines.append(f"• {t}")

    await update.message.reply_text("\n".join(lines))

# ---------- UNLEARN ----------
async def unlearn(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to remove global sources.")
        return

    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="unlearn")

    text = extract_command_text(update)
    rest = strip_command(text, "unlearn")
    if not rest:
        await update.message.reply_text("Usage: /unlearn <exact title shown in /sources>")
        return

    title = rest.strip()
    
    try:
        col = get_collection(chat_id, topic)
    except Exception as e:
        logging.error(f"Failed to get collection: {e}")
        await update.message.reply_text(f"❌ Failed to access database: {e}")
        return
        
    try:
        to_delete = col.get(where={"title": title})
    except Exception as e:
        logging.error(f"Error checking for documents to delete: {e}")
        await update.message.reply_text(f"❌ Error checking for documents: {e}")
        return
        
    removed = len((to_delete or {}).get("ids") or [])

    if removed == 0:
        await update.message.reply_text(
            f"No source titled '{title}' found in topic: {topic} (GLOBAL)."
        )
        return

    try:
        col.delete(where={"title": title})
        await update.message.reply_text(
            f"Removed '{title}' ✅ ({removed} parts) from topic: {topic} (GLOBAL)"
        )
    except Exception as e:
        logging.error(f"Error deleting documents: {e}")
        await update.message.reply_text(f"❌ Failed to remove '{title}': {e}")

# ---------- CLEAR ----------
async def clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ You are not allowed to clear global knowledge.")
        return

    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="clear")

    collection_name = f"{COLLECTION_PREFIX}_{topic}"
    try:
        chroma.delete_collection(collection_name)
        await update.message.reply_text(f"Forgot everything 🧹 (topic: {topic}, scope: GLOBAL)")
    except Exception as e:
        logging.warning(f"Could not delete collection {collection_name}: {e}")
        await update.message.reply_text(
            f"Didn't find any stored sources to clear for topic: {topic} (GLOBAL)."
        )

# ---------- EVALUATION HELPERS ----------
def _eval_context_from_collection(col, extra_query: str = ""):
    eval_docs = []
    qa_docs = []
    try:
        res_eval = col.query(
            query_texts=["evaluation criteria", "guidelines", "rubric", extra_query],
            where={"type": "evaluation"},
            n_results=6,
        )
        eval_docs = res_eval.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying evaluation docs: {e}")
        eval_docs = []

    try:
        res_qa = col.query(
            query_texts=["tips", "examples", "advice", extra_query],
            where={"type": "qa"},
            n_results=6,
        )
        qa_docs = res_qa.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying QA docs: {e}")
        qa_docs = []

    docs = (eval_docs or []) + (qa_docs or [])
    return "\n\n---\n\n".join(docs) if docs else ""

async def run_eval_followup(update: Update, context: ContextTypes.DEFAULT_TYPE, user_request: str):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user

    topic = context.user_data.get("last_eval_topic") or get_current_topic(context)
    pretty_topic = _pretty_topic_for_eval(topic)

    student_text = (context.user_data.get("last_eval_text") or "").strip()
    prior_feedback = (context.user_data.get("last_eval_feedback") or "").strip()

    if not student_text:
        await update.message.reply_text(
            "✅ Evaluation follow-up is ON, but I don't have your last text saved.\n\n"
            "Please paste the full text again (or upload PDF/DOCX), then I'll apply feedback."
        )
        return

    record_event(user.id, topic, kind="eval_followup")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, extra_query=pretty_topic)
    except Exception as e:
        logging.error(f"Error getting collection for eval followup: {e}")
        context_block = ""

    sys_role = _sys_role_for_eval(topic)

    messages = [
        {
            "role": "system",
            "content": (
                sys_role
                + "\n\nRules:\n"
                "- Apply the prior feedback to revise the text.\n"
                "- If the request is vague (e.g., 'do the next step'), default to: "
                "smoother transitions + deeper reflection + stronger conclusion tied to the opening.\n"
                "- Keep the student's voice.\n"
                "- Output ONLY the revised text (no commentary)."
            ),
        },
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {"role": "system", "content": f"Prior feedback to apply:\n{prior_feedback}"},
        {"role": "system", "content": f"Text to revise:\n{student_text}"},
        {"role": "user", "content": f"User request:\n{user_request}"},
    ]

    revised = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.35)

    if revised and not revised.lower().startswith("error"):
        context.user_data["last_eval_text"] = revised
        await send_long(update, revised)
    else:
        await update.message.reply_text("❌ Failed to generate revision. Please try again.")

async def run_eval_qa(update: Update, context: ContextTypes.DEFAULT_TYPE, user_question: str):
    """
    NEW: Answer ANY follow-up question in eval mode using the saved submission + feedback.
    This fixes: "I can't see your essay here" replies.
    """
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user

    topic = context.user_data.get("last_eval_topic") or get_current_topic(context)
    pretty_topic = _pretty_topic_for_eval(topic)

    student_text = (context.user_data.get("last_eval_text") or "").strip()
    prior_feedback = (context.user_data.get("last_eval_feedback") or "").strip()

    if not student_text:
        await update.message.reply_text(
            "✅ Evaluation mode is ON, but I don't have your last submission saved.\n\n"
            "Please paste the full text again (or upload PDF/DOCX)."
        )
        return

    record_event(user.id, topic, kind="eval_qa")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, extra_query=pretty_topic)
    except Exception as e:
        logging.error(f"Error getting collection for eval QA: {e}")
        context_block = ""

    sys = (
        "You are an experienced admissions mentor.\n"
        "The user is asking a follow-up question about a text you already evaluated.\n"
        "Rules:\n"
        "- You DO have access to the student's text below.\n"
        "- Answer the question directly and specifically.\n"
        "- If they ask about grammar, point out concrete issues and show corrected versions of 2–5 short excerpts.\n"
        "- If they ask for suggestions, keep them actionable.\n"
        "- Be concise and readable.\n"
        "- Do NOT say you cannot see the essay.\n"
    )

    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {"role": "system", "content": f"Prior evaluation feedback:\n{prior_feedback}"},
        {"role": "system", "content": f"Student's {pretty_topic}:\n{student_text}"},
        {"role": "user", "content": user_question},
    ]

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.35)
    await send_long(update, a)

async def evaluate_file_for_topic(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="eval")

    allowed = set(EVAL_TOPICS) | {"ielts_writing"}
    if topic not in allowed:
        await update.message.reply_text(
            "To evaluate an essay, recommendation, EC description, portfolio, or IELTS Writing, "
            "choose the correct menu first, then send the file."
        )
        return

    doc = update.message.document
    if not doc:
        await update.message.reply_text("Please attach a PDF or DOCX file.")
        return

    pretty_topic = _pretty_topic_for_eval(topic)
    await update.message.reply_text(f"Reading your {pretty_topic} file…")

    try:
        tgfile = await doc.get_file()
        file_bytes = await tgfile.download_as_bytearray()
        name = (doc.file_name or "document").lower()
    except Exception as e:
        await update.message.reply_text(f"Failed to download file: {e}")
        return

    try:
        if name.endswith(".pdf"):
            full_text = extract_text(io.BytesIO(file_bytes))
        elif name.endswith(".docx"):
            d = DocxDocument(io.BytesIO(file_bytes))
            full_text = "\n".join(p.text for p in d.paragraphs)
        else:
            await update.message.reply_text("Only PDF or DOCX are supported for evaluation.")
            return
    except Exception as e:
        await update.message.reply_text(f"Could not read file: {e}")
        return

    if not (full_text or "").strip():
        await update.message.reply_text("I couldn't read enough text from that file to evaluate.")
        return

    parts = _chunk(full_text, max_chars=1500)
    student_text_for_eval = "\n\n---\n\n".join(parts[:5]) if parts else _truncate_for_storage(full_text, 4000)
    student_text_for_followup = _truncate_for_storage(full_text, 12000)

    await update.message.reply_text(f"Analyzing your {pretty_topic} against my guidelines…")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, extra_query=pretty_topic)
    except Exception as e:
        logging.error(f"Error getting collection for evaluation: {e}")
        context_block = ""

    if topic in {"essays_personal", "essays_supplemental"}:
        sys_role = (
            "You are an expert college admissions essay coach. "
            "Evaluate the student's essay. Focus on clarity, structure, voice, authenticity, and impact. "
            "Give specific, actionable feedback and suggestions."
        )
    elif topic == "recommendations":
        sys_role = (
            "You are an expert on college recommendation letters. "
            "Evaluate the letter in terms of specificity, credibility, depth of insight, and support for the student. "
            "Give constructive feedback and suggestions for improvement."
        )
    elif topic == "extracurriculars":
        sys_role = (
            "You are an expert on extracurricular strategy for college applications. "
            "Evaluate how well the activities are presented in terms of impact, leadership, continuity, and uniqueness. "
            "Give specific, practical suggestions to make the activities stand out."
        )
    elif topic == "ielts_writing":
        sys_role = (
            "You are an experienced IELTS Writing examiner. "
            "Evaluate the student's writing according to IELTS band descriptors. "
            "Comment on Task Response, Coherence and Cohesion, Lexical Resource, and Grammatical Range and Accuracy. "
            "Give an approximate band score and clear, actionable feedback."
        )
    else:
        sys_role = (
            "You are an expert college portfolio reviewer. "
            "Evaluate the portfolio in terms of coherence, originality, technical quality, and fit for selective colleges. "
            "Give specific, constructive feedback, not generic advice."
        )

    messages = [
        {"role": "system", "content": sys_role + " Use the guidelines and examples in the context when relevant."},
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {"role": "user", "content": f"Here is the student's {pretty_topic}. Please evaluate it:\n\n{student_text_for_eval}"},
    ]

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.3)
    set_eval_context(context, topic, student_text_for_followup, a)
    context.user_data["my_eval_count"] = int(context.user_data.get("my_eval_count", 0) or 0) + 1
    await send_long(update, a)

async def evaluate_ielts_writing_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = "ielts_writing"
    record_event(user.id, topic, kind="eval")

    photos = update.message.photo or []
    if not photos:
        await update.message.reply_text("Please send a clear photo of your IELTS Writing answer.")
        return

    largest = photos[-1]
    try:
        tgfile = await largest.get_file()
    except Exception as e:
        await update.message.reply_text(f"Failed to get image file: {e}")
        return
        
    await update.message.reply_text("Reading your IELTS Writing answer from the image…")
    
    try:
        img_bytes = await tgfile.download_as_bytearray()
    except Exception as e:
        await update.message.reply_text(f"Failed to download image: {e}")
        return

    try:
        extracted = extract_text_from_image_bytes(img_bytes)
    except Exception as e:
        await update.message.reply_text(f"Could not extract text from image: {e}")
        return

    if not extracted.strip():
        await update.message.reply_text(
            "I couldn't read enough text from that image. Please try a clearer photo."
        )
        return

    student_text_for_followup = _truncate_for_storage(extracted, 12000)
    parts = _chunk(extracted, max_chars=1500)
    student_text_for_eval = "\n\n---\n\n".join(parts[:5]) if parts else _truncate_for_storage(extracted, 4000)

    await update.message.reply_text("Analyzing your IELTS Writing answer…")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, extra_query="IELTS Writing")
    except Exception as e:
        logging.error(f"Error getting collection for IELTS evaluation: {e}")
        context_block = ""

    sys_role = (
        "You are an experienced IELTS Writing examiner. "
        "Evaluate the student's writing according to IELTS Academic/General Writing band descriptors. "
        "Comment on Task Response, Coherence and Cohesion, Lexical Resource, and Grammatical Range and Accuracy. "
        "Give an approximate band score (like 6.0, 6.5, 7.0) and then clear, actionable feedback."
    )

    messages = [
        {"role": "system", "content": sys_role},
        {"role": "system", "content": f"IELTS writing rubrics and notes (may be empty):\n{context_block}"},
        {"role": "user", "content": f"Here is the student's IELTS Writing answer (from an image):\n\n{student_text_for_eval}"},
    ]

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.3)
    set_eval_context(context, topic, student_text_for_followup, a)
    await send_long(update, a)

async def evaluate_text_for_topic(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_typing(update, context)
    chat_id = update.effective_chat.id
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="eval")

    allowed = set(EVAL_TOPICS) | {"ielts_writing"}
    if topic not in allowed:
        await update.message.reply_text(
            "Text evaluation is only available for Personal Statement, Supplemental Essays, "
            "Extracurriculars, Recommendation Letters, Portfolio, and IELTS Writing.\n\n"
            "Choose the correct topic first, then tap the Evaluation button again."
        )
        return

    raw = (update.message.text or "").strip()
    for marker in [BTN_PS_EVAL, BTN_SUPP_EVAL, BTN_EC_EVAL, BTN_REC_EVAL, BTN_IW_EVAL, BTN_PORT_EVAL]:
        raw = raw.replace(marker, "").strip()

    if len(raw) < 100:
        await update.message.reply_text(
            "The text is too short to evaluate. Please paste the full essay/letter/description."
        )
        return

    pretty_topic = _pretty_topic_for_eval(topic)
    student_text_for_followup = _truncate_for_storage(raw, 12000)
    parts = _chunk(raw, max_chars=1500)
    student_text_for_eval = "\n\n---\n\n".join(parts[:5])

    await update.message.reply_text(f"Evaluating your {pretty_topic}…")

    try:
        col = get_collection(chat_id, topic)
        context_block = _eval_context_from_collection(col, extra_query=pretty_topic)
    except Exception as e:
        logging.error(f"Error getting collection for evaluation: {e}")
        context_block = ""

    if topic in {"essays_personal", "essays_supplemental"}:
        sys_role = (
            "You are an expert college admissions essay coach. "
            "Evaluate the student's essay. Focus on clarity, structure, voice, authenticity, and impact. "
            "Give specific, actionable feedback and suggestions."
        )
    elif topic == "recommendations":
        sys_role = (
            "You are an expert on college recommendation letters. "
            "Evaluate the letter in terms of specificity, credibility, depth of insight, and support for the student. "
            "Give constructive feedback and suggestions for improvement."
        )
    elif topic == "extracurriculars":
        sys_role = (
            "You are an expert on extracurricular strategy for college applications. "
            "Evaluate how well the activities are presented in terms of impact, leadership, continuity, and uniqueness. "
            "Give specific, practical suggestions to make the activities stand out."
        )
    elif topic == "ielts_writing":
        sys_role = (
            "You are an experienced IELTS Writing examiner. "
            "Evaluate the student's writing according to IELTS band descriptors. "
            "Comment on Task Response, Coherence and Cohesion, Lexical Resource, and Grammatical Range and Accuracy. "
            "Give an approximate band score and clear, actionable feedback."
        )
    else:
        sys_role = (
            "You are an expert college portfolio reviewer. "
            "Evaluate the portfolio description in terms of coherence, originality, technical quality, and fit for selective colleges. "
            "Give specific, constructive feedback, not generic advice."
        )

    messages = [
        {"role": "system", "content": sys_role + " Use the guidelines and examples in the context when available."},
        {"role": "system", "content": f"Guidelines + examples (may be empty):\n{context_block}"},
        {"role": "user", "content": f"Here is the student's {pretty_topic}. Please evaluate it:\n\n{student_text_for_eval}"},
    ]

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.3)
    set_eval_context(context, topic, student_text_for_followup, a)
    await send_long(update, a)

# ---------- DOCUMENT & PHOTO ROUTERS ----------
async def document_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    caption = (update.message.caption or "").strip()
    
    # Check for commands in caption FIRST
    if caption:
        if is_caption_command(caption, "teach"):
            await teach(update, context)
            return
        if is_caption_command(caption, "teachrubric"):
            await teachrubric(update, context)
            return
        if is_caption_command(caption, "teachfile_eval"):
            await teachfile_eval(update, context)
            return
        if is_caption_command(caption, "teachfile"):
            await teachfile(update, context)
            return
        if is_caption_command(caption, "teachlink_eval"):
            await teachlink_eval(update, context)
            return
        if is_caption_command(caption, "teachlink"):
            await teachlink(update, context)
            return

    # If no command in caption, check if we're in evaluation mode
    topic = get_current_topic(context)
    if topic in (set(EVAL_TOPICS) | {"ielts_writing"}):
        await evaluate_file_for_topic(update, context)
        return

    await update.message.reply_text(
        "If you want me to LEARN from this file, send it again and write /teachfile, "
        "/teachfile_eval, /teachlink, or /teachlink_eval in the caption.\n\n"
        "If this is an essay, recommendation, EC description, portfolio, or IELTS Writing for feedback, "
        "choose the correct topic and tap its Evaluation button."
    )

async def photo_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    caption = (update.message.caption or "").strip()
    topic = get_current_topic(context)
    user = update.effective_user
    record_event(user.id, topic, kind="photo")

    # Check for commands in caption FIRST
    if caption:
        if is_caption_command(caption, "teach"):
            await teach(update, context)
            return
        if is_caption_command(caption, "teachrubric"):
            await teachrubric(update, context)
            return
        if is_caption_command(caption, "teachimage"):
            await teachimage(update, context)
            return

    if topic == "ielts_writing":
        await evaluate_ielts_writing_image(update, context)
        return

    await update.message.reply_text(
        "I can learn from images too 🤖🖼\n\n"
        "If you want me to *learn* from this image (e.g., essay screenshot, rubric), "
        "send it again with caption:\n\n"
        "/teachimage <title>\n\n"
        "For IELTS Writing evaluation from an image, switch to IELTS Writing first."
    )

# ---------- STATS ----------
async def stats_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    stats = load_stats()
    total_users = len(stats.get("users", []))
    total_msgs = stats.get("messages_total", 0)
    topic_counts = stats.get("topic_counts", {})
    eval_counts = stats.get("eval_counts", {})

    if topic_counts:
        top_topics = sorted(topic_counts.items(), key=lambda kv: kv[1], reverse=True)[:5]
        topics_str = "\n".join(f"- {t}: {c} msgs" for t, c in top_topics)
    else:
        topics_str = "No topic data yet."

    total_evals = sum(eval_counts.values()) if eval_counts else 0

    msg = (
        f"📊 Bot analytics\n"
        f"- Unique users: {total_users}\n"
        f"- Total interactions (events): {total_msgs}\n"
        f"- Total evaluations: {total_evals}\n\n"
        f"Top topics:\n{topics_str}"
    )
    await update.message.reply_text(msg)

# ---------- BACKUP SOURCES ----------
async def backup_sources(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not require_admin(update):
        await update.message.reply_text("⛔ Admin only.")
        return

    await update.message.reply_text("📦 Backing up all sources…")

    data = {}
    try:
        collections = chroma.list_collections()
    except Exception as e:
        await update.message.reply_text(f"Error accessing database: {e}")
        return

    for col_info in collections:
        col_name = col_info.name if hasattr(col_info, "name") else str(col_info)
        if not col_name:
            continue
            
        try:
            col = chroma.get_collection(col_name)
            payload = col.get(include=["documents", "metadatas", "ids"])
            data[col_name] = payload
        except Exception as e:
            logging.error(f"Error backing up collection {col_name}: {e}")
            continue

    path = os.path.join(DATA_DIR, "backup_sources.json")
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        await update.message.reply_text(
            f"✅ Backup completed.\nSaved to:\n{path}\n\nYou can download it from your server volume."
        )
    except Exception as e:
        logging.error(f"Error saving backup: {e}")
        await update.message.reply_text(f"❌ Failed to save backup: {e}")

# ---------- HEALTH CHECK ----------
async def health(update: Update, context: ContextTypes.DEFAULT_TYPE):
    checks = []

    try:
        _ = openai_chat(
            model="gpt-4.1-mini",
            messages=[{"role": "user", "content": "ping"}],
            temperature=0.0,
            max_tokens=5,
        )
        checks.append("✅ OpenAI: OK")
    except Exception as e:
        checks.append(f"❌ OpenAI: {str(e)[:100]}")

    try:
        test_col = chroma.get_or_create_collection("health_check", embedding_function=emb_fn)
        _id = uuid.uuid4().hex
        test_col.add(ids=[_id], documents=["pong"], metadatas=[{"type": "health"}])
        try:
            test_col.delete(ids=[_id])
        except Exception:
            pass
        checks.append("✅ Chroma: writable")
    except Exception as e:
        checks.append(f"❌ Chroma: {str(e)[:100]}")

    try:
        test_path = os.path.join(DATA_DIR, "health.txt")
        with open(test_path, "w", encoding="utf-8") as f:
            f.write("ok")
        os.remove(test_path)
        checks.append(f"✅ Volume: writable ({DATA_DIR})")
    except Exception as e:
        checks.append(f"❌ Volume: {str(e)[:100]}")

    await update.message.reply_text("🧪 Health Check\n\n" + "\n".join(checks))

# ---------- NEW FEATURE HELPERS ----------
def set_pending_feature(context: ContextTypes.DEFAULT_TYPE, feature: str | None):
    context.user_data["pending_feature"] = feature

def clear_pending_feature(context: ContextTypes.DEFAULT_TYPE):
    context.user_data.pop("pending_feature", None)

async def run_brainstorm(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    topic = get_current_topic(context)
    user = update.effective_user
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="brainstorm")

    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)

    try:
        col = get_collection(chat_id, topic)
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying for brainstorm: {e}")
        docs = []
        
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        f"You are an admissions mentor helping a student brainstorm ideas for {nice_topic}.\n"
        "- Give 3-5 short bullet ideas or angles.\n"
        "- Each bullet should be 1-2 concise sentences.\n"
        "- Focus on realistic, personal, and application-relevant ideas.\n"
        "- Keep the tone friendly, specific, and not generic."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {"role": "system", "content": f"Program-specific notes and examples (may be empty):\n{context_block}"}
        )
    messages.append({"role": "user", "content": "Here is the student's situation:\n\n" + description})

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.5)
    await send_long(update, a)

async def brainstorm_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "brainstorm")
        await update.message.reply_text(
            "🧠 Brainstorm mode ON.\n\nBriefly tell me about yourself, your target major, and what you want to write about."
        )
        return
    await run_brainstorm(update, context, parts[1].strip())

async def run_rewrite(update: Update, context: ContextTypes.DEFAULT_TYPE, text_to_fix: str):
    await show_typing(update, context)
    topic = get_current_topic(context)
    user = update.effective_user
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="rewrite")

    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)

    try:
        col = get_collection(chat_id, topic)
        res = col.query(query_texts=[text_to_fix], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying for rewrite: {e}")
        docs = []
        
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        f"You are an admissions writing coach helping improve a student's {nice_topic} text.\n"
        "IMPROVE the text by making it:\n"
        "1. Clearer and more natural\n"
        "2. Better flow and transitions\n"
        "3. Stronger word choice without being fancy\n"
        "4. More impactful for admissions readers\n"
        "5. Preserving the student's original voice and meaning\n\n"
        "OUTPUT FORMAT:\n"
        "📝 IMPROVED VERSION:\n"
        "[The complete rewritten text]\n\n"
        "🔍 KEY IMPROVEMENTS:\n"
        "[Brief point about a key change]\n"
        "[Brief point about another change]\n"
        "[Brief point about the main improvement]\n"
        "(Just 2-3 points, keep it very concise)"
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {"role": "system", "content": f"Program-specific notes and examples (may be empty):\n{context_block}"}
        )
    messages.append({"role": "user", "content": f"Here is the text to improve:\n\n{text_to_fix}"})

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.4)
    await send_long(update, a)

async def rewrite_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "rewrite")
        await update.message.reply_text("✍️ Rewrite mode ON.\n\nSend me the paragraph or essay you want me to improve.")
        return
    await run_rewrite(update, context, parts[1].strip())

async def run_plan(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="plan")

    try:
        col = get_collection(chat_id, topic)
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying for plan: {e}")
        docs = []
        
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        "You are an admissions strategy mentor.\n"
        "- Based on the student's situation, create a concise application plan.\n"
        "- Organize it into short bullet points under 3 headings: Academics & Testing, Essays & Recs, Activities & Extras.\n"
        "- Keep total response around 120-200 words.\n"
        "- Focus on practical next steps, not theory."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append({"role": "system", "content": f"Program-specific planning notes (may be empty):\n{context_block}"})
    messages.append({"role": "user", "content": "Here is my situation:\n\n" + description})

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.5)
    await send_long(update, a)

async def plan_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["topic"] = "application_plan"
    track_topic(context, "application_plan")
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "plan")
        await update.message.reply_text(
            "📅 Application Plan mode ON.\n\nTell me your grade, target countries, intended major, test scores (if any), and your rough deadlines.",
            reply_markup=plan_keyboard(),
        )
        return
    await run_plan(update, context, parts[1].strip())

async def run_recpacket(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="recpacket")

    sys = (
        "You are creating a recommendation letter 'brag sheet' for a teacher.\n"
        "- Output in 3 short sections:\n"
        "  1) 3-5 sentence summary the student can give the teacher.\n"
        "  2) Bullet list of key achievements/impacts.\n"
        "  3) Bullet list of personal qualities and 2-3 specific story ideas.\n"
        "- Keep it concise and realistic for competitive admissions."
    )

    messages = [{"role": "system", "content": sys}, {"role": "user", "content": description}]
    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.5)
    await send_long(update, a)

async def recpacket_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "recpacket")
        await update.message.reply_text(
            "✉️ Rec Letter Packet mode ON.\n\nTell me which teacher will write your rec, what classes you took with them, your achievements, and what you want them to highlight."
        )
        return
    await run_recpacket(update, context, parts[1].strip())

async def run_schoolfinder(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    chat_id = update.effective_chat.id
    record_event(user.id, topic, kind="schoolfinder")

    try:
        col = get_collection(chat_id, topic)
        res = col.query(query_texts=[description], n_results=6)
        docs = res.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying for schoolfinder: {e}")
        docs = []
        
    context_block = "\n\n---\n\n".join(docs) if docs else ""

    sys = (
        "You are a university match advisor.\n"
        "- Based on the student's stats and preferences, suggest Reach, Match, and Safety school types and a few example universities.\n"
        "- For each category, give 2-4 example schools and 1-2 bullets about why they fit.\n"
        "- Keep total response concise (around 150-220 words).\n"
        "- Make it clear this is an approximate starting point and they must research details themselves."
    )

    messages = [{"role": "system", "content": sys}]
    if context_block:
        messages.append(
            {"role": "system", "content": f"Program-specific school lists/notes (may be empty):\n{context_block}"}
        )
    messages.append({"role": "user", "content": description})

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.6)
    await send_long(update, a)

async def schoolfinder_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["topic"] = "school_finder"
    track_topic(context, "school_finder")
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "schoolfinder")
        await update.message.reply_text(
            "🏫 School Finder mode ON.\n\nSend me your GPA, tests, budget, target countries, intended major, and constraints (e.g. need scholarship).",
            reply_markup=schoolfinder_keyboard(),
        )
        return
    await run_schoolfinder(update, context, parts[1].strip())

async def run_portfolioideas(update: Update, context: ContextTypes.DEFAULT_TYPE, description: str):
    await show_typing(update, context)
    user = update.effective_user
    topic = get_current_topic(context)
    record_event(user.id, topic, kind="portfolioideas")

    sys = (
        "You are a portfolio mentor for university applications.\n"
        "- Based on the student's field (e.g. CS, design, art, film, business) and interests, suggest 3-6 concrete project ideas.\n"
        "- Each idea should be 1-2 sentences, focused on impact and what it shows about the student.\n"
        "- Make ideas realistic for a high school student, but impressive."
    )

    messages = [{"role": "system", "content": sys}, {"role": "user", "content": description}]
    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.6)
    await send_long(update, a)

async def portfolioideas_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        set_pending_feature(context, "portfolioideas")
        await update.message.reply_text(
            "🖼️ Portfolio Ideas mode ON.\n\nTell me your field (CS, design, art, film, etc.), your skills, and the kind of programs you are targeting."
        )
        return
    await run_portfolioideas(update, context, parts[1].strip())

# ---------- BOOST TOOLS (CROSS-TOPIC) ----------
def _progress_bar(pct: int) -> str:
    pct = max(0, min(100, int(pct)))
    filled = int(round(pct / 10))
    return "█" * filled + "░" * (10 - filled) + f" {pct}%"

async def tool_my_progress(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Lightweight progress dashboard based on on-bot usage."""
    user = update.effective_user
    stats = load_stats()
    uid = str(user.id)

    msg_count = int(stats.get("messages_per_user", {}).get(uid, 0) or 0)
    evals_done = int(context.user_data.get("my_eval_count", 0) or 0)
    topics_seen = context.user_data.get("topics_seen", []) or []
    tools_used = context.user_data.get("tools_used", []) or []

    engagement = min(100, msg_count * 4)              # 25 messages -> 100
    drafting = min(100, evals_done * 25)              # 4 evals -> 100
    exploration = min(100, len(topics_seen) * 12)     # ~8 topics -> 96
    toolkit = min(100, len(tools_used) * 25)          # 4 tools -> 100

    current = FRIENDLY_TOPIC_NAMES.get(get_current_topic(context), "General")

    out = (
        "📊 MY PROGRESS\n\n"
        f"Engagement:  {_progress_bar(engagement)}\n"
        f"Drafting:    {_progress_bar(drafting)}\n"
        f"Exploration: {_progress_bar(exploration)}\n"
        f"Boost Tools: {_progress_bar(toolkit)}\n\n"
        f"Stats:\n"
        f"• Messages with me: {msg_count}\n"
        f"• Evaluations done: {evals_done}\n"
        f"• Topics explored: {len(topics_seen)}\n"
        f"• Boost tools used: {len(tools_used)}\n"
        f"• Current focus: {current}\n\n"
        "Next step idea: Run 🔍 Find Wow Factor on your latest draft, then apply the suggestions in a rewrite."
    )

    track_tool_use(context, "progress")
    await update.message.reply_text(out, reply_markup=tools_menu_keyboard())

async def tool_insider_tips(update: Update, context: ContextTypes.DEFAULT_TYPE):
    topic = get_current_topic(context)

    tips_by_topic = {
        "essays_personal": (
            "🤫 INSIDER TIPS — PERSONAL STATEMENT\n\n"
            "• Your first 2 lines are everything: start with a moment, not an intro.\n"
            "• One specific scene > 10 generic achievements.\n"
            "• Show reflection: what changed in you, not just what happened.\n"
            "• AOs trust proof: add tiny details (time, place, sensory).\n"
            "• End with forward motion: what you’ll do next in college."
        ),
        "essays_supplemental": (
            "🤫 INSIDER TIPS — SUPPLEMENTALS\n\n"
            "• 'Why us' works best as: YOU → THEIR SPECIFICS → YOU AGAIN.\n"
            "• Mention 2 ultra-specific fit points (lab, prof, program, initiative).\n"
            "• Avoid resume repetition—add new angles and values.\n"
            "• Short prompts: one strong claim + one mini-story + one insight.\n"
            "• Make it sound like a real student, not marketing copy."
        ),
        "extracurriculars": (
            "🤫 INSIDER TIPS — EXTRACURRICULARS\n\n"
            "• Impact beats title. Numbers help (people reached, hours, funds).\n"
            "• Show progression: member → builder → leader/mentor.\n"
            "• Use action verbs + outcomes (what changed because of you).\n"
            "• One 'spike' (a theme) is stronger than 15 random clubs.\n"
            "• Always answer: Why you? Why it matters?"
        ),
        "recommendations": (
            "🤫 INSIDER TIPS — RECOMMENDATIONS\n\n"
            "• Best letters include 2–3 stories that only that teacher could tell.\n"
            "• Ask teachers who saw you struggle AND grow (not just easy A's).\n"
            "• Give them a brag sheet with facts, projects, and specific moments.\n"
            "• Strong recs compare you to peers (“top 5% I’ve taught”).\n"
            "• Remind early: 3 gentle reminders > 1 last-minute panic."
        ),
        "portfolio": (
            "🤫 INSIDER TIPS — PORTFOLIO\n\n"
            "• Curate: 6 great pieces beats 20 average ones.\n"
            "• Add process: drafts, iterations, what you learned.\n"
            "• Label your role clearly (solo vs team, what you owned).\n"
            "• Make it scannable: titles, 1-line context, 1-line takeaway.\n"
            "• Tie to future: what you want to build next."
        ),
        "ielts_writing": (
            "🤫 INSIDER TIPS — IELTS WRITING\n\n"
            "• Task 2: clear position in the intro + topic sentences every paragraph.\n"
            "• Aim for: example → explanation → link back.\n"
            "• Don’t chase fancy words—accuracy > complexity.\n"
            "• Use cohesive devices naturally (however, therefore, moreover).\n"
            "• Save 3 minutes to check grammar + articles + verb tenses."
        ),
    }

    out = tips_by_topic.get(
        topic,
        "🤫 INSIDER TIPS\n\n• Pick a menu topic first for more tailored tips.\n• If you want the fastest improvement: run an evaluation, then ask follow-up questions on the feedback."
    )

    track_tool_use(context, "insider")
    await update.message.reply_text(out, reply_markup=tools_menu_keyboard())

async def tool_power_words(update: Update, context: ContextTypes.DEFAULT_TYPE):
    topic = get_current_topic(context)

    packs = {
        "essays_personal": {
            "Action": ["spearheaded", "orchestrated", "built", "revived", "initiated"],
            "Reflection": ["realized", "reframed", "questioned", "unlearned", "grew into"],
            "Impact": ["shifted", "amplified", "shaped", "opened", "strengthened"],
        },
        "essays_supplemental": {
            "Fit": ["aligned", "intersected", "clicked", "connected", "matched"],
            "Academics": ["inquiry", "research", "seminar", "lab", "capstone"],
            "Community": ["collaborate", "mentor", "contribute", "co-create", "engage"],
        },
        "extracurriculars": {
            "Leadership": ["led", "mobilized", "trained", "scaled", "launched"],
            "Results": ["increased", "reduced", "raised", "delivered", "reached"],
            "Innovation": ["engineered", "prototyped", "automated", "designed", "iterated"],
        },
        "ielts_writing": {
            "Argument": ["therefore", "however", "moreover", "consequently", "nevertheless"],
            "Precision": ["notably", "primarily", "increasingly", "specifically", "ultimately"],
            "Neutral tone": ["suggests", "indicates", "tends to", "is likely to", "can be"],
        },
    }

    pack = packs.get(topic, packs["essays_personal"])
    label = FRIENDLY_TOPIC_NAMES.get(topic, "Your writing")

    lines = [f"⚡ POWER WORDS — {label.upper()}\n"]
    for k, words in pack.items():
        lines.append(f"{k}: " + ", ".join(words))
    lines.append("\nTip: Replace weak verbs (did/helped) with one stronger verb + a result (what changed?).")

    track_tool_use(context, "powerwords")
    await update.message.reply_text("\n".join(lines), reply_markup=tools_menu_keyboard())

async def tool_predict_chances(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """A 'readiness indicator' (NOT a real acceptance probability)."""
    user = update.effective_user
    stats = load_stats()
    uid = str(user.id)

    msg_count = int(stats.get("messages_per_user", {}).get(uid, 0) or 0)
    evals_done = int(context.user_data.get("my_eval_count", 0) or 0)
    topics_seen = context.user_data.get("topics_seen", []) or []
    tools_used = context.user_data.get("tools_used", []) or []

    score = 0
    score += min(35, msg_count)                # up to 35
    score += min(25, evals_done * 8)           # up to 25
    score += min(20, len(topics_seen) * 3)     # up to 20
    score += min(20, len(tools_used) * 4)      # up to 20
    score = max(0, min(100, int(score)))

    if score < 35:
        level = "Early-stage"
        next_steps = "Run 1 evaluation (PS or Supplementals) and use ⚡ Power Words on the revision."
    elif score < 70:
        level = "Building"
        next_steps = "Do 2 evaluations + one rewrite pass focused on clarity + reflection."
    else:
        level = "Strong momentum"
        next_steps = "Polish: tighten openings, add specificity, and align your EC story to your major theme."

    out = (
        "🎯 PREDICT MY CHANCES (READINESS INDICATOR)\n\n"
        f"Readiness score: {_progress_bar(score)}\n"
        f"Status: {level}\n\n"
        "What this is: a rough indicator of how complete/strong your materials are based on your usage here.\n"
        "What this is NOT: a real admissions probability (schools use many factors + external context).\n\n"
        f"Recommended next step: {next_steps}"
    )

    track_tool_use(context, "predict")
    await update.message.reply_text(out, reply_markup=tools_menu_keyboard())

async def run_wowfactor(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str):
    """Analyze a text and highlight the most memorable 'wow factor'."""
    await show_typing(update, context)

    sys = (
        "You are a top college admissions essay coach. "
        "Find the ONE most unique, compelling, memorable element in the student's text — the 'wow factor'.\n\n"
        "Output format (exact headings):\n"
        "🎯 WOW FACTOR: <2-6 word label>\n"
        "✨ Why it stands out: <1-2 sentences>\n"
        "💪 How to amplify it:\n"
        "- <actionable step 1>\n"
        "- <actionable step 2>\n"
        "- <actionable step 3>\n"
        "⚠️ Biggest risk to fix: <1 sentence>\n\n"
        "Be specific. Do not mention AI." 
    )

    content = (text or "").strip()
    content = content[:2500]

    a = openai_chat(
        model="gpt-4.1-mini",
        messages=[{"role": "system", "content": sys}, {"role": "user", "content": content}],
        temperature=0.4,
    )

    track_tool_use(context, "wowfactor")
    await send_long(update, "🔍 FIND WOW FACTOR\n\n" + (a or ""))
    # Keep the tools keyboard visible for the next click.
    await update.message.reply_text("Pick another tool (or tap ⬅️ Back).", reply_markup=tools_menu_keyboard())

# ---------- MAIN ANSWER ----------
async def answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = (update.message.text or "").strip()
    if text.startswith("/"):
        # User sent a command while in a pending mode (rewrite/brainstorm/etc.)
        if context.user_data.get("pending_feature"):
            clear_pending_feature(context)
            await update.message.reply_text(
                "ℹ️ Switched out of the previous mode to process your command."
            )
        return

    logging.info(f"💬 TEXT RECEIVED: {text}")

    chat_id = update.effective_chat.id
    user = update.effective_user
    q = text

    topic_before = get_current_topic(context)
    record_event(user.id, topic_before, kind="message")

    # ---- BACK ----
    if is_back_message(q):
        # Back inside Boost Tools should just exit the tools menu.
        if context.user_data.get("in_tools"):
            context.user_data.pop("in_tools", None)
            clear_pending_feature(context)
            clear_eval_context(context)
            context.user_data["topic"] = DEFAULT_TOPIC
            await update.message.reply_text(
                "Back to main menu.",
                reply_markup=main_menu_keyboard(),
            )
            return

        clear_eval_context(context)
        clear_pending_feature(context)
        context.user_data["topic"] = DEFAULT_TOPIC
        await update.message.reply_text(
            "Back to main menu. You're now in general mode - you can just ask questions or choose a section again.",
            reply_markup=main_menu_keyboard(),
        )
        return

    # quick cancel (optional)
    if q.lower() in {"cancel", "stop", "exit"} and context.user_data.get("eval_active", False):
        clear_eval_context(context)
        await update.message.reply_text("✅ Stopped evaluation follow-up mode.")
        return

    # ---- BOOST TOOLS MENU ----
    if q == BTN_TOOLS:
        clear_eval_context(context)
        clear_pending_feature(context)
        context.user_data["in_tools"] = True
        await update.message.reply_text(
            "🚀 Boost Tools\n\nPick a tool:",
            reply_markup=tools_menu_keyboard(),
        )
        return

    if q in {BTN_PROGRESS, BTN_INSIDER, BTN_POWERWORDS, BTN_PREDICT, BTN_WOWFACTOR}:
        # If the user was in another pending mode, clicking a tool should override it.
        clear_pending_feature(context)
        context.user_data["in_tools"] = True

        if q == BTN_PROGRESS:
            await tool_my_progress(update, context)
            return
        if q == BTN_INSIDER:
            await tool_insider_tips(update, context)
            return
        if q == BTN_POWERWORDS:
            await tool_power_words(update, context)
            return
        if q == BTN_PREDICT:
            await tool_predict_chances(update, context)
            return
        if q == BTN_WOWFACTOR:
            last = (context.user_data.get("last_eval_text") or "").strip()
            if len(last) >= 120:
                await run_wowfactor(update, context, last)
                return

            set_pending_feature(context, "wowfactor")
            await update.message.reply_text(
                "🔍 Find Wow Factor\n\nPaste your essay/paragraph (120+ words).\n"
                "Tip: If you ran an evaluation, I can use your last evaluated text automatically.",
                reply_markup=tools_menu_keyboard(),
            )
            return

    # --- pending feature input ---
    pending = context.user_data.get("pending_feature")
    if pending:
        clear_pending_feature(context)
        if pending == "brainstorm":
            await run_brainstorm(update, context, q)
            return
        if pending == "rewrite":
            await run_rewrite(update, context, q)
            return
        if pending == "plan":
            await run_plan(update, context, q)
            return
        if pending == "recpacket":
            await run_recpacket(update, context, q)
            return
        if pending == "schoolfinder":
            await run_schoolfinder(update, context, q)
            return
        if pending == "portfolioideas":
            await run_portfolioideas(update, context, q)
            return
        if pending == "wowfactor":
            if len((q or "").strip()) < 120:
                await update.message.reply_text(
                    "Please paste a bit more text (120+ words) so I can detect a real wow factor.",
                    reply_markup=tools_menu_keyboard(),
                )
                set_pending_feature(context, "wowfactor")
                return
            await run_wowfactor(update, context, q)
            return

    # ---- MAIN MENUS ----
    if q == BTN_ESSAY:
        clear_eval_context(context)
        await send_with_image(
            update,
            "Essays selected.\nChoose Personal Statement or Supplemental Essays.",
            reply_markup=essay_main_keyboard(),
            image_key="essays_main",
        )
        return

    if q == BTN_SAT:
        clear_eval_context(context)
        await send_with_image(
            update,
            "SAT selected. Choose a section:",
            reply_markup=sat_menu_keyboard(),
            image_key="sat_main",
        )
        return

    if q == BTN_IELTS:
        clear_eval_context(context)
        await send_with_image(
            update,
            "IELTS selected. Choose a section.",
            reply_markup=ielts_main_keyboard(),
            image_key="ielts_main",
        )
        return

    if q == BTN_PLAN_MAIN:
        clear_eval_context(context)
        context.user_data["topic"] = "application_plan"
        track_topic(context, "application_plan")
        set_pending_feature(context, "plan")
        await send_with_image(
            update,
            "📅 Application Plan mode ON.\n\nTell me your grade, target countries, intended major, test scores (if any), and your rough deadlines.",
            reply_markup=plan_keyboard(),
            image_key="plan_main",
        )
        return

    if q == BTN_SF_MAIN:
        clear_eval_context(context)
        context.user_data["topic"] = "school_finder"
        track_topic(context, "school_finder")
        set_pending_feature(context, "schoolfinder")
        await send_with_image(
            update,
            "🏫 School Finder mode ON.\n\nSend me your GPA (or approximate), test scores (if any), budget, target countries, intended major, and constraints (e.g. need scholarship).",
            reply_markup=schoolfinder_keyboard(),
            image_key="schoolfinder_main",
        )
        return

    # ---- TOPIC BUTTONS ----
    if q == BTN_ESSAY_PS:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        await send_with_image(
            update,
            "Great, let's work on your Personal Statement.\nAsk about structure, voice, and storytelling.",
            reply_markup=essay_ps_keyboard(),
            image_key="essays_personal",
        )
        await update.message.reply_text(
            "For detailed feedback, tap '✅ Personal Statement Evaluation' and then upload PDF/DOCX or paste the text.\n\n"
            "After I evaluate, you can ask ANY follow-up question about your essay (grammar, clarity, wording, etc.)."
        )
        return

    if q == BTN_ESSAY_SUPP:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        await send_with_image(
            update,
            "Great, let's work on your Supplemental Essays.\nAsk about 'Why us', community essays, and short prompts.",
            reply_markup=essay_supp_keyboard(),
            image_key="essays_supplemental",
        )
        await update.message.reply_text(
            "For detailed feedback, tap '✅ Supplemental Essay Evaluation' and then upload PDF/DOCX or paste the text.\n\n"
            "After I evaluate, you can ask ANY follow-up question about your essay."
        )
        return

    if q == BTN_EC:
        clear_eval_context(context)
        context.user_data["topic"] = "extracurriculars"
        track_topic(context, "extracurriculars")
        await send_with_image(
            update,
            "Great, let's talk about your Extracurricular activities.\nAsk how to present impact, leadership, and long-term involvement.",
            reply_markup=ec_keyboard(),
            image_key="extracurriculars",
        )
        await update.message.reply_text(
            "For feedback, click '✅ Extracurricular Evaluation' and then upload PDF/DOCX or paste your EC descriptions."
        )
        return

    if q == BTN_EC_PROGRAMS:
        await update.message.reply_text(
            "🌍 Here is a list of top extracurricular programs and opportunities:\n\n"
            "https://docs.google.com/spreadsheets/d/1D-UlJGrg32Ib-9Rvm9y7lKkE6jkx3EK-Kb_qJ6G3tos/edit?usp=sharing\n"
        )
        return

    if q == BTN_REC:
        clear_eval_context(context)
        context.user_data["topic"] = "recommendations"
        track_topic(context, "recommendations")
        await send_with_image(
            update,
            "Great, let's work on Recommendation Letters.\nAsk how to request them and what makes a strong letter.",
            reply_markup=rec_keyboard(),
            image_key="recommendations",
        )
        await update.message.reply_text(
            "You can:\n- Tap '✅ Rec Letter Evaluation' to get feedback on a draft.\n- Tap '📄 Rec Letter Packet' to build a brag sheet."
        )
        return

    if q == BTN_PORT:
        clear_eval_context(context)
        context.user_data["topic"] = "portfolio"
        track_topic(context, "portfolio")
        await send_with_image(
            update,
            "You're now in Portfolio Check.\nAsk about structure and how to present your work.",
            reply_markup=portfolio_keyboard(),
            image_key="portfolio",
        )
        await update.message.reply_text(
            "You can upload PDF/DOCX for detailed feedback (✅ Portfolio Evaluation), or tap '💡 Portfolio Ideas'."
        )
        return

    if q in {BTN_SAT_MATH, BTN_SAT_ENGLISH}:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        nice_name = "SAT Math" if topic == "sat_math" else "SAT English"
        image_key = "sat_math" if topic == "sat_math" else "sat_english"
        await send_with_image(
            update,
            f"You're now in {nice_name}. Ask anything.",
            reply_markup=sat_menu_keyboard(),
            image_key=image_key,
        )
        return

    if q in {BTN_IELTS_READING, BTN_IELTS_LISTENING, BTN_IELTS_SPEAKING}:
        clear_eval_context(context)
        topic = TOPIC_KEYS[q]
        context.user_data["topic"] = topic
        track_topic(context, topic)
        await send_with_image(
            update,
            f"You're now in {FRIENDLY_TOPIC_NAMES.get(topic, topic)}. Ask anything.",
            reply_markup=ielts_main_keyboard(),
            image_key=topic,
        )
        return

    if q == BTN_IELTS_WRITING:
        clear_eval_context(context)
        context.user_data["topic"] = "ielts_writing"
        track_topic(context, "ielts_writing")
        await send_with_image(
            update,
            "You're now in IELTS Writing.\nAsk about Task 1/2, band 7+ strategies, or send your answer for feedback.",
            reply_markup=ielts_writing_keyboard(),
            image_key="ielts_writing",
        )
        await update.message.reply_text(
            "For evaluation, click '✅ Writing Evaluation' then send your answer as text, PDF/DOCX, or a clear photo.\n\n"
            "After evaluation, you can ask ANY follow-up question about your writing."
        )
        return

    # ---- FEATURE BUTTONS ----
    if q == BTN_BRAINSTORM:
        set_pending_feature(context, "brainstorm")
        await update.message.reply_text(
            "🧠 Brainstorm mode ON.\n\nBriefly tell me about yourself, your target major, and what you want to write about."
        )
        return

    if q == BTN_REWRITE:
        set_pending_feature(context, "rewrite")
        await update.message.reply_text("✍️ Rewrite mode ON.\n\nSend me the paragraph or essay you want me to improve.")
        return

    if q == BTN_REC_PACKET:
        set_pending_feature(context, "recpacket")
        await update.message.reply_text(
            "✉️ Rec Letter Packet mode ON.\n\nTell me which teacher will write your rec, what classes you took, your achievements, and what you want highlighted."
        )
        return

    if q == BTN_PORTFOLIO_IDEAS:
        set_pending_feature(context, "portfolioideas")
        await update.message.reply_text(
            "💡 Portfolio Ideas mode ON.\n\nTell me your field (CS, design, art, film, etc.), your skills, and target programs."
        )
        return

    # ---- EVALUATION BUTTONS ----
    if q == BTN_PS_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "essays_personal"
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "essays_personal"
        await update.message.reply_text(
            "Personal Statement Evaluation mode ON ✅\n\nNow paste your Personal Statement (100+ words) or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question about your essay (grammar, clarity, structure, etc.).\n"
            "If you say things like 'apply this feedback' or 'rewrite the conclusion', I will revise the text."
        )
        return

    if q == BTN_SUPP_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "essays_supplemental"
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "essays_supplemental"
        await update.message.reply_text(
            "Supplemental Essay Evaluation mode ON ✅\n\nNow paste your essay (100+ words) or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question about your essay.\n"
            "If you say 'apply this feedback' or 'rewrite the ending', I will revise the text."
        )
        return

    if q == BTN_EC_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "extracurriculars"
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "extracurriculars"
        await update.message.reply_text(
            "Extracurricular Evaluation mode ON ✅\n\nNow paste your EC descriptions or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    if q == BTN_REC_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "recommendations"
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "recommendations"
        await update.message.reply_text(
            "Rec Letter Evaluation mode ON ✅\n\nNow paste the draft letter or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    if q == BTN_IW_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "ielts_writing"
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "ielts_writing"
        await update.message.reply_text(
            "IELTS Writing Evaluation mode ON ✅\n\nSend your answer as text, PDF/DOCX, or a clear photo.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    if q == BTN_PORT_EVAL:
        clear_eval_context(context)
        context.user_data["topic"] = "portfolio"
        context.user_data["eval_active"] = True
        context.user_data["last_eval_topic"] = "portfolio"
        await update.message.reply_text(
            "Portfolio Evaluation mode ON ✅\n\nNow paste your portfolio description or upload PDF/DOCX.\n"
            "After I evaluate, you can ask ANY follow-up question."
        )
        return

    # ---- EVALUATION FLOW (supports follow-ups + Q&A) ----
    if context.user_data.get("eval_active", False):
        last_text = (context.user_data.get("last_eval_text") or "").strip()

        # pasted a new submission while eval mode is ON
        if looks_like_submission(q):
            await evaluate_text_for_topic(update, context)
            return

        # eval mode ON but no submission yet
        if not last_text:
            await update.message.reply_text(
                "✅ Evaluation mode is ON.\n\n"
                "Paste your full text here (100+ words) or upload a PDF/DOCX.\n"
                "After I evaluate, you can ask ANY follow-up question about your text."
            )
            return

        # ✅ If it looks like a rewrite/apply request -> do rewrite
        if is_followup_intent(q):
            await run_eval_followup(update, context, q)
            return

        # ✅ Otherwise: answer ANY follow-up question using saved text + feedback
        await update.message.reply_text("Got it! Thinking about your question…")
        await run_eval_qa(update, context, q)
        return

    # ---- NORMAL Q&A WITH RAG ----
    await show_typing(update, context)
    await update.message.reply_text("Got it! Thinking about your question…")

    topic = get_current_topic(context)
    
    try:
        col = get_collection(chat_id, topic)
        results = col.query(query_texts=[q], where={"type": "qa"}, n_results=4)
        docs = results.get("documents", [[]])[0]
    except Exception as e:
        logging.error(f"Error querying collection: {e}")
        docs = []

    if not docs:
        try:
            results = col.query(query_texts=[q], n_results=4)
            docs = results.get("documents", [[]])[0]
        except Exception as e:
            logging.error(f"Error querying collection (fallback): {e}")
            docs = []

    context_block = "\n\n---\n\n".join(docs or [])
    nice_topic = FRIENDLY_TOPIC_NAMES.get(topic, topic)

    sys = (
        "You are UniVenture, a focused university admissions mentor.\n"
        f"Current topic: {nice_topic}.\n"
        "- Give short, clear, high-value answers (3-7 sentences max).\n"
        "- Speak in a natural, human tone, like a friendly but direct older student mentor.\n"
        "- Prioritize practical, actionable advice over theory.\n"
        "- Use the provided context (if any) as trusted program material and do not contradict it."
    )

    messages = [
        {"role": "system", "content": sys},
        {"role": "system", "content": f"Context (may be empty):\n{context_block}"},
        {"role": "user", "content": q},
    ]

    a = openai_chat(model="gpt-4.1-mini", messages=messages, temperature=0.4)
    await send_long(update, a)

# -------- Dummy HTTP server for Render/Railway --------
def start_dummy_server():
    class Handler(BaseHTTPRequestHandler):
        def do_GET(self):
            self.send_response(200)
            self.send_header("Content-type", "text/plain")
            self.end_headers()
            self.wfile.write(b"OK")

        def log_message(self, format, *args):
            return

    port = int(os.environ.get("PORT", 10000))
    server = HTTPServer(("0.0.0.0", port), Handler)
    print(f"✅ Health check server running on port {port}")
    server.serve_forever()

# ===== Setup Application =====
app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()

# ===== GROUP 0: COMMANDS ONLY =====
app.add_handler(CommandHandler("start", start), group=0)
app.add_handler(CommandHandler("teach", teach), group=0)
app.add_handler(CommandHandler("teachrubric", teachrubric), group=0)
app.add_handler(CommandHandler("teachfile", teachfile), group=0)
app.add_handler(CommandHandler("teachfile_eval", teachfile_eval), group=0)
app.add_handler(CommandHandler("teachlink", teachlink), group=0)
app.add_handler(CommandHandler("teachlink_eval", teachlink_eval), group=0)
app.add_handler(CommandHandler("teachimage", teachimage), group=0)
app.add_handler(CommandHandler("sources", sources), group=0)
app.add_handler(CommandHandler("sources_all", sources_all), group=0)
app.add_handler(CommandHandler("unlearn", unlearn), group=0)
app.add_handler(CommandHandler("clear", clear), group=0)
app.add_handler(CommandHandler("stats", stats_cmd), group=0)
app.add_handler(CommandHandler("brainstorm", brainstorm_cmd), group=0)
app.add_handler(CommandHandler("rewrite", rewrite_cmd), group=0)
app.add_handler(CommandHandler("plan", plan_cmd), group=0)
app.add_handler(CommandHandler("recpacket", recpacket_cmd), group=0)
app.add_handler(CommandHandler("schoolfinder", schoolfinder_cmd), group=0)
app.add_handler(CommandHandler("portfolioideas", portfolioideas_cmd), group=0)
app.add_handler(CommandHandler("backup_sources", backup_sources), group=0)
app.add_handler(CommandHandler("health", health), group=0)

# ===== GROUP 1: FILES / PHOTOS =====
app.add_handler(MessageHandler(filters.Document.ALL, document_router), group=1)
app.add_handler(MessageHandler(filters.PHOTO, photo_router), group=1)

# ===== GROUP 2: NORMAL TEXT (ABSOLUTELY LAST) =====
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, answer), group=2)

print(
    "\n" + "="*80 + "\n"
    "Bot is running with GLOBAL per-topic RAG + metadata separation (qa/evaluation) + submenus "
    "+ eval follow-ups (apply feedback) + eval Q&A (any follow-up question) + embedded tools "
    "+ Application Plan & School Finder + analytics + admin locks + backup + health "
    "+ UUID IDs + robust command parsing (incl captions)…\n"
    "="*80 + "\n"
)

if __name__ == "__main__":
    threading.Thread(target=start_dummy_server, daemon=True).start()
    print("✅ Bot starting...")
    app.run_polling()
